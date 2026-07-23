"""Collect sweep results into a complete, organised analysis package.

Reads the per-run result caches the pipeline writes (_method_{profile}_wy{Y}_
c2e{F}.pkl and _original_wy{Y}_results.pkl in OUTPUT_ROOT). Works on whatever
has finished so far (re-runnable mid-sweep). Produces, in
OUTPUT_ROOT/sweep_results_<stamp>/:

  README.md                    what every file is
  SWEEP_SUMMARY.md             key numbers in prose
  tables/
    sweep_master.csv           one row per (weather_year, future), ~35 metrics
                               + original-run references + deltas
    stats_by_future.csv        mean/median/min/max/std of key metrics
    design_ranking.csv         designs ranked robust -> fragile, per future
    shed_by_region_long.csv    region shedding, tidy long format
    shed_by_region_wide_{F}.csv  region x weather-year matrix per future (GWh)
    shed_by_sector_long.csv    heat vs electricity vs other, per run
    unserved_by_month_long.csv seasonality of unserved energy, per run
    worst_events.csv           top shortfall events of every run
  figures/
    fig01..fig10 (see FIGURES below)
  report/
    SWEEP_REPORT.docx          the crafted cross-sweep report: headline stats,
                               all figures with captions, fragile/robust
                               design tables, regional and event tables,
                               provenance and caveats
"""
from __future__ import annotations

import os as _os, sys as _sys
_sys.path.insert(0, _os.path.dirname(_os.path.abspath(__file__)))
import os
import re
import glob
import time
import pickle
import argparse

import numpy as np
import pandas as pd
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt

import config
import systems

FUT_COLORS = ['#1C7293', '#C0682E', '#2E7D5B', '#7A4E8C']


# ---------------------------------------------------------------------------
# loading
# ---------------------------------------------------------------------------
def load_runs(profile: str):
    """[(system, design_id, weather_year|None, future, mod, orig|None), ...]
    sorted. Legacy gotske caches ('_method_mixed_wy2015_c2e2042.pkl') and
    v21 multi-system caches ('..._{system}--{design}_c2e{F}.pkl') both load."""
    pat = os.path.join(config.OUTPUT_ROOT, f"_method_{profile}_*_c2e*.pkl")
    rx = re.compile(r"_method_" + re.escape(profile) + r"_(?P<key>.+)_c2e(?P<f>\d{4})\.pkl$")
    runs = []
    for p in sorted(glob.glob(pat)):
        m = rx.search(os.path.basename(p))
        if not m:
            continue
        key, f = m.group('key'), int(m.group('f'))
        sysname, did = systems.split_design_key(key)
        y = int(did[2:]) if (sysname == 'gotske' and did.startswith('wy')) else None
        try:
            with open(p, 'rb') as fh:
                mod = pickle.load(fh)
        except Exception as e:
            print(f"  skip unreadable {os.path.basename(p)}: {e}")
            continue
        org = None
        op = os.path.join(config.OUTPUT_ROOT, f"_original_{key}_results.pkl")
        if os.path.exists(op):
            try:
                with open(op, 'rb') as fh:
                    org = pickle.load(fh)
            except Exception:
                org = None
        runs.append((sysname, did, y, f, mod, org))
    runs.sort(key=lambda t: (t[0], t[3], t[2] if t[2] is not None else -1, t[1]))
    return runs


def _g(d, k, default=np.nan):
    try:
        v = d.get(k, default)
        return float(v) if np.isscalar(v) else default
    except Exception:
        return default


def master_table(runs) -> pd.DataFrame:
    rows = []
    for sysname, did, y, f, mod, org in runs:
        r = {'system': sysname, 'design_id': did, 'weather_year': y,
             'c2e_future': f}
        r['objective_BEUR'] = _g(mod, 'objective') / 1e9
        for k, name in [
            ('total_demand_TWh', 'demand_TWh'),
            ('load_shedding_TWh', 'unserved_total_TWh'),
            ('unserved_elec_TWh', 'unserved_elec_TWh'),
            ('unserved_heat_TWh', 'unserved_heat_TWh'),
            ('resource_adequacy_pct', 'adequacy_pct'),
            ('unserved_energy_pct', 'unserved_pct'),
            ('gotske_resource_adequacy_pct', 'gotske_adequacy_pct'),
            ('gotske_unserved_elec_TWh', 'gotske_unserved_elec_TWh'),
            ('elec_demand_TWh', 'elec_demand_TWh'),
            ('peak_unserved_MW', 'peak_unserved_MW'),
            ('n_shortfall_events', 'n_events'),
            ('n_elec_events_over_24h', 'n_elec_events_over_24h'),
            ('max_event_duration_h', 'max_event_h'),
            ('mean_event_duration_h', 'mean_event_h'),
            ('hours_with_shortfall', 'shortfall_hours'),
            ('co2_emissions_Mt', 'co2_Mt'),
            ('co2_emissions_pct_of_1990', 'co2_pct_of_1990'),
            ('backup_energy_TWh', 'backup_TWh'),
            ('peak_backup_MW', 'peak_backup_MW'),
            ('curtailment_TWh', 'curtailment_TWh'),
            ('re_share_pct', 're_share_pct'),
            ('co2_price_EUR_per_t', 'co2_price'),
        ]:
            r[name] = _g(mod, k)
        ev = (mod.get('worst_events') or [{}])[0]
        r['worst_event_energy_GWh'] = _g(ev, 'energy_GWh') if ev else np.nan
        r['worst_event_duration_h'] = _g(ev, 'duration_h') if ev else np.nan
        try:
            r['worst_event_month'] = pd.Timestamp(ev['start']).month if ev else np.nan
        except Exception:
            r['worst_event_month'] = np.nan
        r['pipeline_version'] = str(mod.get('pipeline_version', ''))
        if org:
            r['orig_unserved_TWh'] = _g(org, 'load_shedding_TWh')
            r['orig_co2_Mt'] = _g(org, 'co2_emissions_Mt')
            r['orig_objective_BEUR'] = _g(org, 'objective') / 1e9
            r['orig_curtailment_TWh'] = _g(org, 'curtailment_TWh')
            r['orig_backup_TWh'] = _g(org, 'backup_energy_TWh')
            r['orig_re_share_pct'] = _g(org, 're_share_pct')
            r['delta_unserved_TWh'] = r['unserved_total_TWh'] - r['orig_unserved_TWh']
            r['delta_co2_Mt'] = r['co2_Mt'] - r['orig_co2_Mt']
            r['delta_curtailment_TWh'] = r['curtailment_TWh'] - r['orig_curtailment_TWh']
        rows.append(r)
    return (pd.DataFrame(rows)
            .sort_values(['system', 'c2e_future', 'weather_year', 'design_id'])
            .reset_index(drop=True))


def detail_tables(runs, tdir: str):
    """Region / sector / month / event tables from the per-run dicts."""
    reg, sec, mon, evs = [], [], [], []
    for sysname, did, y, f, mod, _ in runs:
        for k, v in (mod.get('load_shedding_by_region_GWh') or {}).items():
            reg.append({'system': sysname, 'design_id': did, 'weather_year': y, 'c2e_future': f, 'region': k, 'shed_GWh': float(v)})
        for k, v in (mod.get('load_shedding_by_sector_GWh') or {}).items():
            sec.append({'system': sysname, 'design_id': did, 'weather_year': y, 'c2e_future': f, 'sector': k, 'shed_GWh': float(v)})
        for m, v in (mod.get('unserved_by_month_GWh') or {}).items():
            mon.append({'system': sysname, 'design_id': did, 'weather_year': y, 'c2e_future': f, 'month': int(m), 'unserved_GWh': float(v)})
        for rank, ev in enumerate(mod.get('worst_events') or [], 1):
            evs.append({'system': sysname, 'design_id': did, 'weather_year': y, 'c2e_future': f, 'rank': rank, **ev})
    regdf = pd.DataFrame(reg); secdf = pd.DataFrame(sec)
    mondf = pd.DataFrame(mon); evdf = pd.DataFrame(evs)
    if len(regdf):
        regdf.to_csv(os.path.join(tdir, 'shed_by_region_long.csv'), index=False)
        for f in sorted(regdf['c2e_future'].unique()):
            w = (regdf[regdf['c2e_future'] == f]
                 .pivot_table(index='region', columns='weather_year',
                              values='shed_GWh', aggfunc='sum', fill_value=0.0))
            w = w.loc[w.sum(axis=1).sort_values(ascending=False).index]
            w.round(1).to_csv(os.path.join(tdir, f'shed_by_region_wide_{f}.csv'))
    if len(secdf):
        secdf.to_csv(os.path.join(tdir, 'shed_by_sector_long.csv'), index=False)
    if len(mondf):
        mondf.to_csv(os.path.join(tdir, 'unserved_by_month_long.csv'), index=False)
    if len(evdf):
        evdf.to_csv(os.path.join(tdir, 'worst_events.csv'), index=False)
    return regdf, secdf, mondf, evdf


def stats_and_ranking(df: pd.DataFrame, tdir: str):
    keys = ['unserved_total_TWh', 'unserved_elec_TWh', 'unserved_heat_TWh',
            'adequacy_pct', 'gotske_adequacy_pct', 'co2_Mt', 'backup_TWh',
            'curtailment_TWh', 'max_event_h', 'peak_unserved_MW']
    stats = (df.groupby('c2e_future')[[k for k in keys if k in df.columns]]
               .agg(['mean', 'median', 'min', 'max', 'std']).round(3))
    stats.columns = ['_'.join(c) for c in stats.columns]   # flat, readable headers
    stats = stats.T                                        # metrics as rows
    stats.index.name = 'metric'
    stats.to_csv(os.path.join(tdir, 'stats_by_future.csv'))
    piv = df.pivot_table(index='design_id', columns='c2e_future',
                         values='unserved_total_TWh')
    rank = piv.rank(ascending=True)  # 1 = most robust (least unserved)
    rank.columns = [f'rank_{c}' for c in rank.columns]
    out = pd.concat([piv.add_prefix('unserved_TWh_'), rank], axis=1)
    out['mean_rank'] = rank.mean(axis=1)
    out = out.sort_values('mean_rank')
    out.round(3).to_csv(os.path.join(tdir, 'design_ranking.csv'))
    return stats, out


# ---------------------------------------------------------------------------
# figures
# ---------------------------------------------------------------------------
def make_figures(df, regdf, mondf, evdf, runs, fdir: str):
    df_y = df[df['weather_year'].notna()]          # year-axis (gotske) subset
    futs = sorted(df['c2e_future'].unique())
    C = {f: FUT_COLORS[i % len(FUT_COLORS)] for i, f in enumerate(futs)}
    made = []

    def save(name):
        p = os.path.join(fdir, name)
        plt.tight_layout(); plt.savefig(p, dpi=150); plt.close(); made.append(name)

    # 01 unserved vs design year
    plt.figure(figsize=(9.5, 4.6))
    for f in futs:
        d = df_y[df_y['c2e_future'] == f]
        plt.plot(d['weather_year'], d['unserved_total_TWh'], 'o-', color=C[f],
                 lw=1.7, ms=4.5, label=f'C2E {f} total')
        plt.plot(d['weather_year'], d['unserved_elec_TWh'], 's--', color=C[f],
                 lw=1.1, ms=3, alpha=.55, label=f'C2E {f} electricity')
    plt.xlabel('design weather year'); plt.ylabel('unserved energy (TWh)')
    plt.title('Unserved energy of each fixed design under future climate')
    plt.legend(fontsize=8); plt.grid(alpha=.25)
    save('fig01_unserved_by_design_year.png')

    # 02 exceedance across designs
    plt.figure(figsize=(6.6, 4.6))
    for f in futs:
        v = np.sort(df.loc[df['c2e_future'] == f, 'unserved_total_TWh'].dropna().values)
        if len(v):
            plt.step(v, 1 - (np.arange(len(v)) + .5) / len(v), where='post',
                     color=C[f], label=f'C2E {f}')
    plt.xlabel('unserved energy (TWh)'); plt.ylabel('fraction of designs at or above')
    plt.title('Adequacy tail across designs'); plt.legend(); plt.grid(alpha=.25)
    save('fig02_adequacy_exceedance.png')

    # 03 CO2 drift
    plt.figure(figsize=(9.5, 4.6))
    for f in futs:
        d = df_y[df_y['c2e_future'] == f]
        plt.plot(d['weather_year'], d['co2_Mt'], 'o-', color=C[f], lw=1.7, ms=4.5,
                 label=f'C2E {f}')
    if 'orig_co2_Mt' in df_y.columns:
        d0 = df_y.drop_duplicates('weather_year')
        plt.plot(d0['weather_year'], d0['orig_co2_Mt'], 'k.:', lw=1,
                 label='original (own weather)')
    plt.axhline(0, color='k', lw=.6)
    plt.xlabel('design weather year'); plt.ylabel('net CO2 (Mt)')
    plt.title('Drift from carbon neutrality under future climate')
    plt.legend(fontsize=8); plt.grid(alpha=.25)
    save('fig03_co2_drift.png')

    # 04 shed vs emit plane
    plt.figure(figsize=(6.6, 5.2))
    for f in futs:
        d = df[df['c2e_future'] == f]
        plt.scatter(d['co2_Mt'], d['unserved_total_TWh'], s=30, color=C[f],
                    label=f'C2E {f}')
        for _, r in d.iterrows():
            _lab = (str(int(r['weather_year'])) if pd.notna(r['weather_year'])
                    else str(r['design_id']).split('__')[0])
            plt.annotate(_lab, (r['co2_Mt'], r['unserved_total_TWh']),
                         fontsize=6, alpha=.6, xytext=(3, 3), textcoords='offset points')
    plt.xlabel('net CO2 (Mt)  [emit response]')
    plt.ylabel('unserved energy (TWh)  [shed response]')
    plt.title('How each design pays for future weather')
    plt.legend(); plt.grid(alpha=.25)
    save('fig04_shed_vs_emit.png')

    # 05 region x year heatmaps
    if len(regdf):
        for f in futs:
            w = (regdf[regdf['c2e_future'] == f]
                 .pivot_table(index='region', columns='weather_year',
                              values='shed_GWh', aggfunc='sum', fill_value=0.0))
            if w.empty:
                continue
            w = w.loc[w.sum(axis=1).sort_values(ascending=False).index].head(20)
            plt.figure(figsize=(max(7, .45 * len(w.columns) + 3), .34 * len(w) + 2.2))
            plt.imshow(np.log10(w.values + 1), aspect='auto', cmap='YlOrRd')
            plt.colorbar(label='log10(shed GWh + 1)')
            plt.yticks(range(len(w.index)), w.index, fontsize=8)
            plt.xticks(range(len(w.columns)), w.columns, rotation=60, fontsize=8)
            plt.title(f'Load shedding by region and design year, C2E {f}')
            save(f'fig05_heatmap_region_year_{f}.png')

    # 06 persistence: 2042 vs 2099 per design
    if len(futs) >= 2:
        piv = df.pivot_table(index='weather_year', columns='c2e_future',
                             values='unserved_total_TWh').dropna()
        if len(piv) >= 2:
            f1, f2 = futs[0], futs[1]
            plt.figure(figsize=(5.8, 5.4))
            plt.scatter(piv[f1], piv[f2], s=34, color='#444')
            for y, r in piv.iterrows():
                plt.annotate(int(y), (r[f1], r[f2]), fontsize=6, alpha=.6,
                             xytext=(3, 3), textcoords='offset points')
            lim = [0, max(piv.max()) * 1.06 + 1e-6]
            plt.plot(lim, lim, 'k--', lw=.8, alpha=.6)
            plt.xlim(lim); plt.ylim(lim)
            plt.xlabel(f'unserved under C2E {f1} (TWh)')
            plt.ylabel(f'unserved under C2E {f2} (TWh)')
            plt.title('Does design vulnerability persist across futures?')
            plt.grid(alpha=.25)
            save('fig06_scatter_futures.png')

    # 07 monthly profile of unserved energy
    if len(mondf):
        plt.figure(figsize=(8.4, 4.4))
        for f in futs:
            d = mondf[mondf['c2e_future'] == f]
            if d.empty:
                continue
            g = d.groupby('month')['unserved_GWh']
            m = g.mean().reindex(range(1, 13), fill_value=0)
            lo = g.min().reindex(range(1, 13), fill_value=0)
            hi = g.max().reindex(range(1, 13), fill_value=0)
            plt.plot(m.index, m.values, 'o-', color=C[f], label=f'C2E {f} mean')
            plt.fill_between(m.index, lo.values, hi.values, color=C[f], alpha=.15)
        plt.xticks(range(1, 13), ['J', 'F', 'M', 'A', 'M', 'J', 'J', 'A', 'S', 'O', 'N', 'D'])
        plt.xlabel('month'); plt.ylabel('unserved energy (GWh)')
        plt.title('When the stress happens: monthly unserved energy across designs '
                  '(band = min-max over designs)')
        plt.legend(); plt.grid(alpha=.25)
        save('fig07_unserved_monthly_profile.png')

    # 08 worst-event timing
    if len(evdf):
        top = evdf[evdf['rank'] == 1].copy()
        if len(top):
            top['month'] = pd.to_datetime(top['start'], errors='coerce').dt.month
            plt.figure(figsize=(8.2, 4.6))
            for f in futs:
                d = top[top['c2e_future'] == f]
                plt.scatter(d['month'], d['duration_h'],
                            s=np.sqrt(d['energy_GWh'].clip(lower=1)) * 3,
                            color=C[f], alpha=.75, label=f'C2E {f}')
            plt.xticks(range(1, 13), ['J', 'F', 'M', 'A', 'M', 'J', 'J', 'A', 'S', 'O', 'N', 'D'])
            plt.xlabel('month of worst event'); plt.ylabel('duration (h)')
            plt.title('Worst shortfall event of each run: timing and duration '
                      '(bubble = energy)')
            plt.legend(); plt.grid(alpha=.25)
            save('fig08_worst_event_timing.png')

    # 09 heat vs electricity stacked bars
    plt.figure(figsize=(10, 4.6))
    n = len(futs); width = .8 / max(n, 1)
    yrs = sorted(df['weather_year'].unique())
    x = np.arange(len(yrs))
    for i, f in enumerate(futs):
        d = df[df['c2e_future'] == f].set_index('weather_year').reindex(yrs)
        he = d['unserved_heat_TWh'].fillna(0).values
        el = d['unserved_elec_TWh'].fillna(0).values
        plt.bar(x + i * width, he, width, color=C[f], alpha=.85,
                label=f'C2E {f} heat')
        plt.bar(x + i * width, el, width, bottom=he, color=C[f], alpha=.4,
                label=f'C2E {f} electricity')
    plt.xticks(x + width * (n - 1) / 2, yrs, rotation=60, fontsize=8)
    plt.ylabel('unserved energy (TWh)'); plt.xlabel('design weather year')
    plt.title('What is unserved: heat vs electricity')
    plt.legend(fontsize=8, ncol=2); plt.grid(axis='y', alpha=.25)
    save('fig09_heat_vs_elec.png')

    # 10 unserved duration curves overlay
    have_dc = any(isinstance(mod.get('_unserved_duration_curve'), pd.Series)
                  for _, _, _, _, mod, _ in runs)
    if have_dc:
        plt.figure(figsize=(8.2, 4.6))
        for sysname, did, y, f, mod, _ in runs:
            dc = mod.get('_unserved_duration_curve')
            if isinstance(dc, pd.Series) and len(dc):
                plt.plot(np.arange(len(dc)), dc.values / 1e3, color=C[f],
                         alpha=.35, lw=.9)
        for f in futs:
            plt.plot([], [], color=C[f], label=f'C2E {f}')
        plt.xlabel('snapshots, sorted'); plt.ylabel('unserved power (GW)')
        plt.title('Unserved-power duration curves, one line per run')
        plt.legend(); plt.grid(alpha=.25)
        plt.xlim(0, None)
        save('fig10_duration_curves.png')

    # 11: per-design bars when non-gotske rows exist (no year axis there)
    df_d = df[df['weather_year'].isna()]
    if len(df_d):
        for sysname in sorted(df_d['system'].unique()):
            dd = df_d[df_d['system'] == sysname]
            piv = dd.pivot_table(index='design_id', columns='c2e_future',
                                 values='unserved_total_TWh')
            piv.index = [i.replace('__', '\n') for i in piv.index]
            ax = piv.plot.barh(figsize=(9, max(3, .5 * len(piv))),
                               color=[C[f] for f in piv.columns])
            ax.set_xlabel('unserved energy (TWh)')
            ax.set_title(f'Unserved energy per design: {sysname}')
            ax.grid(axis='x', alpha=.25)
            save(f'fig11_unserved_by_design_{sysname}.png')
    return made


# ---------------------------------------------------------------------------
# report + docs
# ---------------------------------------------------------------------------
FIG_CAPTIONS = {
    'fig01': 'Unserved energy of every fixed design (x: the weather year it was designed on) when operated under each C2E future. Solid: total; dashed: electricity only.',
    'fig02': 'Exceedance view of the same data: the fraction of designs at or above a given unserved-energy level. The right-hand tail is the robustness story.',
    'fig03': 'Net CO2 of each design under future weather, against its own-weather original (dotted). Distance from zero is the drift from designed carbon neutrality.',
    'fig04': 'The two ways a fixed system can pay for unplanned weather: emitting more (x) or serving less (y). Each point is one design.',
    'fig05': 'Where the shedding happens: load shedding by region (rows, top 20) and design year (columns), log colour scale.',
    'fig06': 'Vulnerability persistence: each design plotted by its unserved energy under the two futures. Points near the diagonal fail (or hold) consistently.',
    'fig07': 'When the stress happens: monthly unserved energy, mean across designs with min-max band.',
    'fig08': 'The single worst shortfall event of every run: its calendar month, duration, and energy (bubble size).',
    'fig09': 'Composition of unserved energy: heat (solid) vs electricity (pale) per design and future.',
    'fig10': 'Unserved-power duration curves of every run: how deep and how long shortfalls run.',
}


def build_docx(df, stats, ranking, regdf, evdf, made, outdir, profile, title):
    try:
        from docx import Document
        from docx.shared import Inches
    except Exception as e:
        print(f"  (SWEEP_REPORT.docx skipped: python-docx unavailable: {e})")
        return None

    def tbl(doc, frame, caption, max_rows=25):
        if frame is None or not len(frame):
            return
        doc.add_paragraph(caption).runs[0].bold = True
        f2 = frame.head(max_rows).reset_index() if frame.index.name or \
            not isinstance(frame.index, pd.RangeIndex) else frame.head(max_rows)
        t = doc.add_table(rows=1, cols=len(f2.columns))
        t.style = 'Light Grid Accent 1'
        for j, c in enumerate(f2.columns):
            t.rows[0].cells[j].text = str(c)
        for _, row in f2.iterrows():
            cells = t.add_row().cells
            for j, v in enumerate(row):
                cells[j].text = ('' if pd.isna(v) else
                                 f'{v:,.3f}' if isinstance(v, float) else str(v))

    doc = Document()
    doc.add_heading(title, 0)
    futs = sorted(df['c2e_future'].unique())
    yrs = sorted(df['weather_year'].unique())
    doc.add_paragraph(
        f"Sweep of {len(yrs)} fixed Gotske designs (weather years {yrs[0]}-{yrs[-1]}) "
        f"re-dispatched under C2E futures {', '.join(map(str, futs))}. Profile: {profile}. "
        f"Scenario: {getattr(config, 'SCENARIO_LABEL', '')}. Pipeline "
        f"{sorted(set(df['pipeline_version'].astype(str)))}. {len(df)} completed runs "
        f"collected {time.strftime('%Y-%m-%d %H:%M')}.").italic = True

    doc.add_heading('1. Headline statistics', 1)
    for f in futs:
        d = df[df['c2e_future'] == f]
        doc.add_paragraph(
            f"C2E {f}: unserved energy mean {d['unserved_total_TWh'].mean():.2f} TWh "
            f"(median {d['unserved_total_TWh'].median():.2f}, "
            f"worst {d['unserved_total_TWh'].max():.2f} TWh at design "
            f"wy{int(d.loc[d['unserved_total_TWh'].idxmax(), 'weather_year']) if d['unserved_total_TWh'].notna().any() else '?'}); "
            f"electricity-only mean {d['unserved_elec_TWh'].mean():.3f} TWh; "
            f"net CO2 mean {d['co2_Mt'].mean():.1f} Mt (max {d['co2_Mt'].max():.1f}); "
            f"curtailment mean {d['curtailment_TWh'].mean():.0f} TWh; "
            f"backup mean {d['backup_TWh'].mean():.0f} TWh.")
    tbl(doc, stats, 'Table 1. Distribution of key metrics per future (across designs).')

    doc.add_heading('2. Robust and fragile designs', 1)
    tbl(doc, ranking.head(5), 'Table 2a. Five most ROBUST designs (lowest unserved, mean rank across futures).')
    tbl(doc, ranking.tail(5).iloc[::-1], 'Table 2b. Five most FRAGILE designs.')

    doc.add_heading('3. Figures', 1)
    for name in made:
        p = os.path.join(outdir, 'figures', name)
        if os.path.exists(p):
            doc.add_picture(p, width=Inches(6.3))
            cap = FIG_CAPTIONS.get(name.split('_')[0], '')
            doc.add_paragraph(f"{name}: {cap}").italic = True

    if len(regdf):
        doc.add_heading('4. Regional concentration', 1)
        agg = (regdf.groupby(['c2e_future', 'region'])['shed_GWh'].sum()
               .reset_index().sort_values('shed_GWh', ascending=False))
        tbl(doc, agg.head(15), 'Table 3. Total load shedding by region, summed over designs (GWh).')

    if len(evdf):
        doc.add_heading('5. Worst events across the sweep', 1)
        top = evdf.sort_values('energy_GWh', ascending=False).head(10)[
            ['weather_year', 'c2e_future', 'start', 'duration_h', 'peak_MW', 'energy_GWh']]
        tbl(doc, top, 'Table 4. Ten largest shortfall events anywhere in the sweep.')

    doc.add_heading('6. Provenance and caveats', 1)
    for t in [
        "Each run freezes one published design's capacities and re-solves operation only "
        "under the stated C2E future; the original (own-weather) dispatch of the same "
        "design is the reference. All runs share solver settings, carbon price and "
        "load-shedding cost, so gaps are attributable to the imposed climate.",
        "Each C2E year file is a SINGLE YEAR of one realization: differences between "
        "futures, and between designs under one future, mix the forced climate signal "
        "with internal variability. Ensemble-mean signals from the C2E reference "
        "publication are not directly comparable to these single-year results.",
        "Per-channel imposition: DIRECT for wind, solar, run-of-river and cooling; "
        "QDM (annual-conserving change factor) for heating demand and reservoir inflow. "
        "Direct supply carries the C2E-vs-network conversion level offset, visible "
        "mainly as curtailment; interpret curtailment changes accordingly.",
        "Per-run detail (full report, channel tables, event tables, dispatched networks) "
        "lives in each run's own folder; this document is the cross-sweep view.",
    ]:
        doc.add_paragraph(t)

    path = os.path.join(outdir, 'report', 'SWEEP_REPORT.docx')
    doc.save(path)
    return path


README = """# Sweep results package

Produced by `python -m src.collect_sweep` on whatever runs had finished.
Re-run it any time; it makes a fresh timestamped folder.

- SWEEP_SUMMARY.md ......... key numbers in prose
- report/SWEEP_REPORT.docx . the crafted cross-sweep report (stats, all
                             figures with captions, robust/fragile designs,
                             regional + event tables, provenance)
- tables/sweep_master.csv .. ONE ROW PER RUN (weather_year x future) with all
                             headline metrics, original-run references, deltas
- tables/stats_by_future.csv distribution of key metrics across designs
- tables/design_ranking.csv  designs ranked robust -> fragile per future
- tables/shed_by_region_*.csv regional shedding (long + per-future matrices)
- tables/shed_by_sector_long.csv  heat vs electricity vs other
- tables/unserved_by_month_long.csv  seasonality per run
- tables/worst_events.csv .. top shortfall events of every run
- figures/fig01..fig10 ..... see captions inside the report

Per-run deliverables (full REPORT docx, channel signal tables, per-run
figures, dispatched networks) are in each run's own folder under
output/<sweep-name>/wyYYYY_fFFFF/.
"""


def summary_md(df, outdir, profile):
    futs = sorted(df['c2e_future'].unique())
    _sys = ", ".join(sorted(df['system'].unique()))
    lines = [f"# Sweep summary ({profile})", f"Systems: {_sys}", "",
             f"Runs collected: {len(df)} "
             f"({df['weather_year'].nunique()} weather years x "
             f"{[int(x) for x in futs]})", ""]
    for f in futs:
        d = df[df['c2e_future'] == f]
        lines += [f"## C2E {f}",
                  f"- unserved energy: mean {d['unserved_total_TWh'].mean():.2f} TWh, "
                  f"median {d['unserved_total_TWh'].median():.2f}, "
                  f"worst {d['unserved_total_TWh'].max():.2f} TWh",
                  f"- electricity-only unserved: mean {d['unserved_elec_TWh'].mean():.3f} TWh "
                  f"(Gotske-basis adequacy min {d['gotske_adequacy_pct'].min():.3f}%)"
                  if 'gotske_adequacy_pct' in d else "",
                  f"- net CO2: mean {d['co2_Mt'].mean():.1f} Mt, max {d['co2_Mt'].max():.1f} Mt",
                  f"- curtailment: mean {d['curtailment_TWh'].mean():.0f} TWh; "
                  f"backup: mean {d['backup_TWh'].mean():.0f} TWh", ""]
    lines += ["Caveat: single-year C2E slices mix forced signal with internal "
              "variability; see report section 6."]
    with open(os.path.join(outdir, 'SWEEP_SUMMARY.md'), 'w') as fh:
        fh.write("\n".join(l for l in lines if l is not None))


def main(argv=None):
    ap = argparse.ArgumentParser()
    ap.add_argument('--profile', default=getattr(config, 'RUN_PROFILE', 'mixed'))
    ap.add_argument('--outdir', default=None)
    ap.add_argument('--title', default='Climate-robustness sweep: fixed designs under future weather')
    args = ap.parse_args(argv)

    runs = load_runs(args.profile)
    if not runs:
        print("no completed runs found for profile", args.profile)
        return 1
    outdir = args.outdir or os.path.join(
        config.OUTPUT_ROOT, 'sweep_results_' + time.strftime('%Y%m%d_%H%M'))
    for sub in ('tables', 'figures', 'report'):
        os.makedirs(os.path.join(outdir, sub), exist_ok=True)

    df = master_table(runs)
    df.to_csv(os.path.join(outdir, 'tables', 'sweep_master.csv'), index=False)
    regdf, secdf, mondf, evdf = detail_tables(runs, os.path.join(outdir, 'tables'))
    stats, ranking = stats_and_ranking(df, os.path.join(outdir, 'tables'))
    made = make_figures(df, regdf, mondf, evdf, runs, os.path.join(outdir, 'figures'))
    summary_md(df, outdir, args.profile)
    with open(os.path.join(outdir, 'README.md'), 'w') as fh:
        fh.write(README)
    rpt = build_docx(df, stats, ranking, regdf, evdf, made, outdir,
                     args.profile, args.title)

    print(f"collected {len(df)} runs -> {outdir}")
    print(f"  tables : {len(os.listdir(os.path.join(outdir, 'tables')))} files")
    print(f"  figures: {len(made)}")
    if rpt:
        print(f"  report : {rpt}")
    cols = ['weather_year', 'c2e_future', 'unserved_total_TWh',
            'unserved_elec_TWh', 'co2_Mt', 'adequacy_pct']
    print(df[[c for c in cols if c in df.columns]].to_string(index=False))
    return 0


if __name__ == '__main__':
    raise SystemExit(main())
