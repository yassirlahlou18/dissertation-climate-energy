"""Builds the figures, tables and Word reports. See the pipeline guide PDF."""

from __future__ import annotations

import os as _os, sys as _sys
_sys.path.insert(0, _os.path.dirname(_os.path.abspath(__file__)))
import os
import numpy as np
import pandas as pd
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

import config

BLUE = '#457b9d'; RED = '#e76f51'; GREEN = '#2a9d8f'; AMBER = '#f39c12'; GREY = '#888888'
SUBDIRS = ['01_generation', '02_load_shedding', '03_curtailment', '04_regional',
           '05_sector_coupling', '06_storage', '07_climate_signal',
           '08_reliability_deep', 'tables', 'networks']


def _dirs(method):
    base = config.run_dir(method)
    paths = {'base': base}
    for s in SUBDIRS:
        d = os.path.join(base, s); os.makedirs(d, exist_ok=True); paths[s] = d
    return paths


def _save(path):
    plt.tight_layout(); plt.savefig(path, dpi=150, bbox_inches='tight'); plt.close()


def _bar(series, title, ylab, path, color=BLUE, horizontal=False, rot=40):
    series = series.dropna()
    if series.empty:
        plt.close(); return False
    if horizontal:
        plt.figure(figsize=(9, max(4, len(series) * 0.28)))
        series.sort_values().plot.barh(color=color, edgecolor='k', lw=.3)
        plt.xlabel(ylab)
    else:
        plt.figure(figsize=(max(8, len(series) * 0.5), 5))
        series.plot.bar(color=color, edgecolor='k', lw=.3)
        plt.ylabel(ylab); plt.xticks(rotation=rot, ha='right')
    plt.title(title); plt.grid(axis='x' if horizontal else 'y', alpha=.3)
    _save(path); return True


# ----------------------------------------------------------------------------
# figures + tables for one solved network
# ----------------------------------------------------------------------------
def make_figures(results, method, label, cf_change_df=None, suffix=''):
    P = _dirs(method); ts = results['_timeseries']; w = ts['weights']
    figs = {}

    def T(name, sub):  # table path; suffix distinguishes original-run companions
        if suffix and name.endswith('.csv'):
            name = name[:-4] + suffix + '.csv'
        return os.path.join(P['tables'], name)

    def F(folder, name):  # figure path with suffix
        if suffix and name.endswith('.png'):
            name = name[:-4] + suffix + '.png'
        return os.path.join(P[folder], name)

    # ---------- 01 generation ----------
    gen = pd.Series(results['gen_by_carrier_TWh']).drop('load_shedding', errors='ignore')
    gen = gen[gen > 0.01].sort_values(ascending=False)
    gen.to_csv(T('gen_by_carrier.csv', None), header=['TWh'])
    p = F('01_generation', 'genmix.png')
    if _bar(gen, f'Generation by technology - {label}', 'Annual generation (TWh)', p):
        figs['genmix'] = p

    gbc = ts['gen_by_carrier']; annual = gbc.multiply(w, axis=0).sum()
    majors = [c for c in annual.index if annual[c] > annual.sum() * 0.005 and c != 'load_shedding']
    if majors:
        weekly = gbc[majors].resample('7D').mean()
        plt.figure(figsize=(15, 5)); weekly.plot.area(alpha=.85, lw=.4, ax=plt.gca())
        plt.ylabel('Average power (MW)'); plt.title(f'Weekly generation - {label}')
        plt.legend(fontsize=7, ncol=2, loc='upper right'); plt.grid(alpha=.3)
        p = F('01_generation', 'weekly_stack.png'); _save(p); figs['weekly'] = p

        monthly = gbc.multiply(w, axis=0).resample('ME').sum() / 1e6
        monthly[majors].to_csv(T('gen_monthly.csv', None))
        plt.figure(figsize=(13, 5))
        monthly[majors].plot.bar(stacked=True, alpha=.85, edgecolor='white', lw=.3, ax=plt.gca())
        plt.ylabel('Generation (TWh)'); plt.title(f'Monthly generation - {label}')
        plt.legend(fontsize=7, ncol=3, loc='upper right'); plt.grid(axis='y', alpha=.3)
        plt.xticks(range(len(monthly)), [d.strftime('%b') for d in monthly.index], rotation=0)
        p = F('01_generation', 'monthly_stack.png'); _save(p); figs['monthly'] = p

    # RE share + utilisation
    util = pd.Series(results.get('re_utilisation_by_carrier', {}))
    if not util.empty:
        util.to_csv(T('re_utilisation.csv', None), header=['utilisation'])
        p = F('01_generation', 're_utilisation.png')
        _bar(util, f'RE capacity utilisation - {label}', 'Achieved / potential', p, color=GREEN)

    # ---------- 02 load shedding ----------
    ls = ts['load_shedding']
    if ls.max() > 0:
        plt.figure(figsize=(15, 3.8))
        plt.plot(ls.index, ls.values, color='#e74c3c', lw=.5)
        plt.fill_between(ls.index, ls.values, alpha=.3, color='#e74c3c')
        plt.ylabel('Load shedding (MW)'); plt.title(f'Load shedding over time - {label}'); plt.grid(alpha=.3)
        p = F('02_load_shedding', 'timeseries.png'); _save(p); figs['shedding_ts'] = p
        # monthly shed
        ls_m = (ls * w).resample('ME').sum() / 1e6
        ls_m.to_csv(T('load_shedding_monthly.csv', None), header=['TWh'])
        plt.figure(figsize=(11, 4))
        ls_m.plot.bar(color='#e74c3c', edgecolor='k', lw=.3, ax=plt.gca())
        plt.ylabel('Load shedding (TWh)'); plt.title(f'Monthly load shedding - {label}')
        plt.xticks(range(len(ls_m)), [d.strftime('%b') for d in ls_m.index], rotation=0); plt.grid(axis='y', alpha=.3)
        p = F('02_load_shedding', 'monthly.png'); _save(p); figs['shedding_monthly'] = p

    for key, fname, color in [('load_shedding_by_bus_GWh', 'by_bus', '#e74c3c'),
                              ('load_shedding_by_region_GWh', 'by_region', '#c0392b'),
                              ('load_shedding_by_sector_GWh', 'by_sector', '#922b21')]:
        s = pd.Series(results.get(key, {}))
        if not s.empty:
            s.to_csv(T(f'load_shedding_{fname}.csv', None), header=['GWh'])
            p = os.path.join(P['02_load_shedding'], f'{fname}.png')
            _bar(s, f'Load shedding {fname.replace("_", " ")} - {label}', 'GWh', p, color=color, horizontal=True)
            figs[f'shedding_{fname}'] = p
    # sector shed time series
    if 'load_shedding_by_sector' in ts and ts['load_shedding_by_sector'].abs().max().max() > 0:
        sshed = ts['load_shedding_by_sector'].resample('7D').mean()
        plt.figure(figsize=(15, 4)); sshed.plot.area(alpha=.8, lw=.3, ax=plt.gca())
        plt.ylabel('Load shedding (MW)'); plt.title(f'Load shedding by sector (weekly) - {label}')
        plt.legend(fontsize=7, ncol=2); plt.grid(alpha=.3)
        p = F('02_load_shedding', 'by_sector_timeseries.png'); _save(p)

    # ---------- 03 curtailment ----------
    curt = pd.Series(results.get('curt_by_carrier_TWh', {})).sort_values(ascending=False)
    if not curt.empty:
        curt.to_csv(T('curtailment_by_carrier.csv', None), header=['TWh'])
        p = F('03_curtailment', 'by_carrier.png')
        _bar(curt, f'Curtailment by technology - {label}', 'Curtailment (TWh)', p, color=AMBER)
        figs['curtailment'] = p
    rate = pd.Series(results.get('curtailment_rate_by_carrier', {})) * 100
    if not rate.empty:
        rate.to_csv(T('curtailment_rate.csv', None), header=['pct'])
        p = F('03_curtailment', 'rate_by_carrier.png')
        _bar(rate, f'Curtailment rate by technology - {label}', '% of potential curtailed', p, color='#d68910')
    cregion = pd.Series(results.get('curt_by_region_TWh', {}))
    if not cregion.empty:
        cregion.to_csv(T('curtailment_by_region.csv', None), header=['TWh'])
        p = F('03_curtailment', 'by_region.png')
        _bar(cregion, f'Curtailment by region - {label}', 'Curtailment (TWh)', p, color=AMBER, horizontal=True)

    # ---------- 04 regional ----------
    reg = pd.Series(results['gen_by_region_TWh']).sort_values()
    reg.to_csv(T('gen_by_region.csv', None), header=['TWh'])
    p = F('04_regional', 'gen_by_region.png')
    _bar(reg, f'Electricity generation by region - {label}', 'Annual generation (TWh)', p, horizontal=True)
    figs['region'] = p
    # region x carrier heatmap-style table
    rc = results.get('gen_by_region_carrier_TWh', {})
    if rc:
        rows = {}
        for k, v in rc.items():
            r_, c_ = k.split('|', 1)
            rows.setdefault(r_, {})[c_] = v
        rc_df = pd.DataFrame(rows).T.fillna(0)
        rc_df.to_csv(T('gen_by_region_carrier.csv', None))
        # stacked bar of major carriers by region
        major_c = [c for c in rc_df.columns if rc_df[c].sum() > rc_df.values.sum() * 0.01
                   and c != 'load_shedding']
        if major_c:
            plt.figure(figsize=(13, 6))
            rc_df[major_c].plot.bar(stacked=True, ax=plt.gca(), edgecolor='white', lw=.2, alpha=.9)
            plt.ylabel('Generation (TWh)'); plt.title(f'Generation by region and technology - {label}')
            plt.legend(fontsize=7, ncol=3); plt.xticks(rotation=60, ha='right'); plt.grid(axis='y', alpha=.3)
            p = F('04_regional', 'region_x_technology.png'); _save(p); figs['region_tech'] = p
    demreg = pd.Series(results.get('demand_by_region_TWh', {}))
    if not demreg.empty:
        demreg.to_csv(T('demand_by_region.csv', None), header=['TWh'])

    # ---------- 05 sector coupling ----------
    le = pd.Series(results.get('link_energy_TWh', {})).head(20)
    if not le.empty:
        le.to_csv(T('link_energy.csv', None), header=['TWh'])
        p = F('05_sector_coupling', 'link_energy.png')
        _bar(le, f'Sector-coupling link energy - {label}', 'Throughput (TWh)', p, color='#6c5ce7', horizontal=True)
    dem_sec = pd.Series(results.get('demand_by_sector_TWh', {}))
    if not dem_sec.empty:
        dem_sec.to_csv(T('demand_by_sector.csv', None), header=['TWh'])
        p = F('05_sector_coupling', 'demand_by_sector.png')
        _bar(dem_sec, f'Demand by sector - {label}', 'Demand (TWh)', p, color='#0984e3', horizontal=True)
    lf = pd.Series(results.get('line_flows_TWh', {})).head(20)
    if not lf.empty:
        lf.to_csv(T('transmission_flows.csv', None), header=['TWh'])
        p = F('05_sector_coupling', 'transmission_flows.png')
        _bar(lf, f'Top transmission flows - {label}', 'Energy transmitted (TWh)', p, color='#00b894', horizontal=True)

    # supply-demand balance
    if ts['total_demand'].sum() > 0:
        bal = pd.DataFrame({'Generation': ts['total_gen'], 'Demand': ts['total_demand']}).resample('7D').mean() / 1e3
        plt.figure(figsize=(15, 4.5))
        plt.plot(bal.index, bal['Generation'], color=GREEN, lw=1.4, label='Generation')
        plt.plot(bal.index, bal['Demand'], color=RED, lw=1.4, label='Demand')
        plt.fill_between(bal.index, bal['Generation'], bal['Demand'], where=bal['Generation'] >= bal['Demand'], alpha=.15, color=GREEN)
        plt.fill_between(bal.index, bal['Generation'], bal['Demand'], where=bal['Generation'] < bal['Demand'], alpha=.15, color=RED)
        plt.ylabel('Power (GW, weekly avg)'); plt.title(f'Supply-demand balance - {label}'); plt.legend(); plt.grid(alpha=.3)
        p = F('05_sector_coupling', 'supply_demand_balance.png'); _save(p); figs['balance'] = p

    # ---------- 06 storage ----------
    if 'store_e' in ts:
        se = ts['store_e']; big = [c for c in se.columns if se[c].max() > 0.1]
        if big:
            plt.figure(figsize=(15, 4.5)); se[big].plot(lw=1, ax=plt.gca())
            plt.ylabel('Energy stored (TWh)'); plt.title(f'Store state of charge - {label}')
            plt.legend(fontsize=7, ncol=2); plt.grid(alpha=.3)
            p = F('06_storage', 'store_soc.png'); _save(p); figs['store_soc'] = p
            (se[big]).to_csv(T('store_soc.csv', None))
    sd = pd.Series(results.get('storage_discharge_TWh', {}))
    if not sd.empty:
        sd.to_csv(T('storage_discharge.csv', None), header=['TWh'])

    # ---------- 07 climate signal ----------
    if cf_change_df is not None and not cf_change_df.empty:
        cf_change_df.to_csv(T('cf_change_detail.csv', None), index=False)
        bycar = cf_change_df.groupby('carrier').agg(
            mean_pct=('pct_change_mean', 'mean'), std_pct=('pct_change_std', 'mean')).reset_index()
        bycar.to_csv(T('cf_change_by_carrier.csv', None), index=False)
        fig, axes = plt.subplots(1, 2, figsize=(13, 5))
        axes[0].bar(bycar['carrier'], bycar['mean_pct'], color=BLUE, edgecolor='k', lw=.3)
        axes[0].set_title('Mean CF change by technology'); axes[0].axhline(0, color='k', lw=.6)
        axes[0].set_ylabel('% change in mean CF')
        axes[1].bar(bycar['carrier'], bycar['std_pct'], color=RED, edgecolor='k', lw=.3)
        axes[1].set_title('Variability (std) change by technology'); axes[1].axhline(0, color='k', lw=.6)
        axes[1].set_ylabel('% change in CF std')
        for ax in axes:
            ax.tick_params(axis='x', rotation=30); ax.grid(axis='y', alpha=.3)
        p = F('07_climate_signal', 'cf_change_by_technology.png'); _save(p); figs['cf_signal'] = p
        # by region
        byreg = cf_change_df.groupby('region').agg(mean_pct=('pct_change_mean', 'mean')).reset_index()
        byreg.to_csv(T('cf_change_by_region.csv', None), index=False)
        p = F('07_climate_signal', 'cf_change_by_region.png')
        _bar(byreg.set_index('region')['mean_pct'], 'Mean CF change by region',
             '% change in mean CF', p, color=BLUE, horizontal=True)

    # ---- 08 reliability deep (Gotske-style: loss of load, events, emissions) ----
    D = P['08_reliability_deep']
    # loss-of-load time series + cumulative unserved
    if '_unserved_ts' in results:
        uts = results['_unserved_ts']
        uts.to_csv(T('unserved_energy_timeseries_MW.csv', None), header=['MW'])
        if float(uts.max()) > 0:
            fig, ax = plt.subplots(figsize=(13, 3.2))
            ax.fill_between(range(len(uts)), uts.values, color=RED, linewidth=0)
            ax.set_title(f'Loss of load over time: {label}', fontweight='bold')
            ax.set_xlabel('snapshot (3-hourly)'); ax.set_ylabel('Unserved power (MW)')
            _save(os.path.join(D, 'loss_of_load_timeseries.png')); figs['lol_ts'] = os.path.join(D, 'loss_of_load_timeseries.png')
    if '_unserved_cumulative_TWh' in results:
        cum = results['_unserved_cumulative_TWh']
        cum.to_csv(T('unserved_energy_cumulative_TWh.csv', None), header=['TWh'])
        if float(cum.iloc[-1]) > 0:
            fig, ax = plt.subplots(figsize=(11, 3.4))
            ax.plot(range(len(cum)), cum.values, color=RED, linewidth=1.5)
            ax.set_title(f'Cumulative unserved energy: {label}', fontweight='bold')
            ax.set_xlabel('snapshot (3-hourly)'); ax.set_ylabel('Cumulative unserved (TWh)')
            _save(os.path.join(D, 'unserved_cumulative.png')); figs['lol_cum'] = os.path.join(D, 'unserved_cumulative.png')
    # monthly unserved
    if results.get('unserved_by_month_GWh'):
        ubm = pd.Series(results['unserved_by_month_GWh']).sort_index()
        ubm.index = ['Jan','Feb','Mar','Apr','May','Jun','Jul','Aug','Sep','Oct','Nov','Dec']
        ubm.to_csv(T('unserved_by_month_GWh.csv', None), header=['GWh'])
        if ubm.sum() > 0:
            _bar(ubm, f'Unserved energy by month: {label}', 'GWh',
                 os.path.join(D, 'unserved_by_month.png'), color=RED)
    # emissions by technology + by month
    if results.get('co2_emissions_by_tech_Mt'):
        ebt = pd.Series(results['co2_emissions_by_tech_Mt']).sort_values(ascending=False)
        ebt.to_csv(T('emissions_by_technology_Mt.csv', None), header=['Mt'])
        _bar(ebt, f'CO2 emissions by technology: {label}', 'Mt CO2',
             os.path.join(D, 'emissions_by_technology.png'), color=BLUE, horizontal=True)
    if results.get('co2_emissions_by_month_Mt'):
        ebm = pd.Series(results['co2_emissions_by_month_Mt']).sort_index()
        ebm.index = ['Jan','Feb','Mar','Apr','May','Jun','Jul','Aug','Sep','Oct','Nov','Dec']
        ebm.to_csv(T('emissions_by_month_Mt.csv', None), header=['Mt'])
    # backup activation profile
    if '_backup_ts' in results:
        bts = results['_backup_ts']
        bts.to_csv(T('backup_generation_timeseries_MW.csv', None), header=['MW'])
        if float(bts.max()) > 0:
            fig, ax = plt.subplots(figsize=(13, 3.2))
            ax.fill_between(range(len(bts)), bts.values, color='#8d6e63', linewidth=0)
            ax.set_title(f'Backup (emitting) generation over time: {label}', fontweight='bold')
            ax.set_xlabel('snapshot (3-hourly)'); ax.set_ylabel('Backup power (MW)')
            _save(os.path.join(D, 'backup_activation_timeseries.png')); figs['backup_ts'] = os.path.join(D, 'backup_activation_timeseries.png')
    # loss-of-load duration curve
    if '_unserved_duration_curve' in results:
        dc = results['_unserved_duration_curve']
        dc.to_csv(T('unserved_duration_curve_MW.csv', None), header=['MW'])
        if float(dc.max()) > 0:
            fig, ax = plt.subplots(figsize=(8, 3.4))
            ax.plot(np.arange(len(dc)) / len(dc) * 100, dc.values, color=RED, linewidth=1.5)
            ax.set_title(f'Loss-of-load duration curve: {label}', fontweight='bold')
            ax.set_xlabel('% of time unserved power exceeds level'); ax.set_ylabel('Unserved power (MW)')
            _save(os.path.join(D, 'unserved_duration_curve.png')); figs['lol_dc'] = os.path.join(D, 'unserved_duration_curve.png')
    # worst shortfall events table
    if results.get('worst_events'):
        we = pd.DataFrame(results['worst_events'])
        we.to_csv(T('worst_shortfall_events.csv', None), index=False)
        figs['_worst_events'] = we

    # ---- Gotske adequacy table (electricity-only headline + heat separately) ----
    adeq = pd.DataFrame([
        {'metric': 'Electricity demand, exogenous + endogenous (TWh)', 'value': results.get('elec_demand_TWh', float('nan'))},
        {'metric': 'Unserved ELECTRICITY (TWh) [Gotske headline]', 'value': results.get('gotske_unserved_elec_TWh', results.get('unserved_elec_TWh', 0))},
        {'metric': 'Resource adequacy, electricity (%) [Gotske]', 'value': results.get('gotske_resource_adequacy_pct', float('nan'))},
        {'metric': 'Unserved HEAT (TWh)', 'value': results.get('unserved_heat_TWh', 0)},
        {'metric': 'Unserved total, electricity + heat (TWh)', 'value': results.get('load_shedding_TWh', 0)},
        {'metric': 'Electricity shedding events longer than 24 h', 'value': results.get('n_elec_events_over_24h', 0)},
    ])
    adeq.to_csv(T('adequacy_gotske.csv', None), index=False)
    figs['_adeq'] = adeq

    # ---- cost decomposition (makes the objective interpretable: the raw
    # objective mixes operations, carbon payments and shedding penalties) ----
    try:
        price = results.get('co2_price_EUR_per_t', None)
        em = results.get('co2_emissions_Mt', None)
        shed_cost_B = results.get('load_shedding_TWh', 0.0) * 1e6 * float(
            getattr(config, 'LOAD_SHEDDING_COST', 1e5)) / 1e9
        carbon_B = (price * em * 1e6 / 1e9) if (price is not None and em is not None) else float('nan')
        obj_B = results.get('objective', float('nan')) / 1e9
        other_B = obj_B - shed_cost_B - (carbon_B if carbon_B == carbon_B else 0.0)
        costd = pd.DataFrame([
            {'component': 'Objective total (B EUR)', 'value': obj_B},
            {'component': 'Load shedding penalty (B EUR)', 'value': shed_cost_B},
            {'component': 'Net carbon payment, price x net emissions (B EUR)', 'value': carbon_B},
            {'component': 'Other operating cost, residual (B EUR)', 'value': other_B},
        ])
        costd.to_csv(T('cost_decomposition.csv', None), index=False)
        figs['_costd'] = costd
    except Exception:
        pass

    # ---- seasonal summary + stress indicators ----
    try:
        VRE = ['solar', 'solar rooftop', 'onwind', 'offwind-ac', 'offwind-dc', 'ror']
        gbc_ts = ts['gen_by_carrier']
        vre_cols = [c for c in VRE if c in gbc_ts.columns]
        vre_ts = gbc_ts[vre_cols].sum(axis=1) if vre_cols else pd.Series(0.0, index=gbc_ts.index)
        dem_ser = ts['total_demand']
        bts_ser = results.get('_backup_ts', pd.Series(0.0, index=dem_ser.index))
        em_ser = results.get('_emissions_ts', pd.Series(0.0, index=dem_ser.index))
        uns_ser = results.get('_unserved_ts', ts.get('load_shedding', pd.Series(0.0, index=dem_ser.index)))
        if not isinstance(uns_ser, pd.Series):
            uns_ser = pd.Series(0.0, index=dem_ser.index)
        wser = pd.Series(w.values, index=dem_ser.index)
        seas_map = {12: 'DJF', 1: 'DJF', 2: 'DJF', 3: 'MAM', 4: 'MAM', 5: 'MAM',
                    6: 'JJA', 7: 'JJA', 8: 'JJA', 9: 'SON', 10: 'SON', 11: 'SON'}
        seas = pd.Series([seas_map[m] for m in dem_ser.index.month], index=dem_ser.index)
        rows = []
        for sname in ['DJF', 'MAM', 'JJA', 'SON']:
            m = seas == sname
            rows.append({'season': sname,
                         'demand_TWh': float((dem_ser[m] * wser[m]).sum() / 1e6),
                         'vre_gen_TWh': float((vre_ts[m] * wser[m]).sum() / 1e6),
                         'backup_TWh': float((bts_ser[m] * wser[m]).sum() / 1e6),
                         'unserved_GWh': float((uns_ser[m] * wser[m]).sum() / 1e3),
                         'emissions_Mt': float((em_ser[m] * wser[m]).sum() / 1e6)})
        pd.DataFrame(rows).to_csv(T('seasonal_summary.csv', None), index=False)

        residual = (dem_ser - vre_ts)
        roll3d = residual.rolling('3D').mean() if hasattr(residual, 'rolling') else residual
        weekly_vre_share = (vre_ts.resample('7D').mean() /
                            dem_ser.resample('7D').mean().replace(0, np.nan)) * 100
        stress = pd.DataFrame([
            {'indicator': 'Peak residual load, demand minus VRE (GW)', 'value': float(residual.max()) / 1e3},
            {'indicator': 'Max 3-day mean residual load (GW)', 'value': float(roll3d.max()) / 1e3},
            {'indicator': 'Min weekly VRE share of demand (%)', 'value': float(weekly_vre_share.min())},
            {'indicator': 'Peak backup (emitting) power (GW)', 'value': float(bts_ser.max()) / 1e3},
            {'indicator': 'Peak unserved power (GW)', 'value': float(uns_ser.max()) / 1e3},
        ])
        stress.to_csv(T('stress_indicators.csv', None), index=False)
        figs['_stress'] = stress
        # residual load duration curve figure
        rsorted = np.sort(residual.values)[::-1] / 1e3
        fig, ax = plt.subplots(figsize=(8, 3.4))
        ax.plot(np.arange(len(rsorted)) / len(rsorted) * 100, rsorted, color=BLUE, lw=1.5)
        ax.set_title(f'Residual load duration curve: {label}', fontweight='bold')
        ax.set_xlabel('% of time residual load exceeds level'); ax.set_ylabel('Residual load (GW)')
        ax.axhline(0, color='grey', lw=.6)
        _save(F('08_reliability_deep', 'residual_load_duration.png'))
        figs['residual_dc'] = F('08_reliability_deep', 'residual_load_duration.png')
    except Exception:
        pass

    # the scalar deep-metrics table (the heart of the reliability analysis)
    deep = pd.DataFrame([
        {'metric': 'Resource adequacy (%)', 'value': results.get('resource_adequacy_pct', float('nan'))},
        {'metric': 'Unserved energy (% of demand)', 'value': results.get('unserved_energy_pct', 0)},
        {'metric': 'Unserved energy (TWh)', 'value': results.get('load_shedding_TWh', 0)},
        {'metric': 'Peak unserved power (MW)', 'value': results.get('peak_unserved_MW', 0)},
        {'metric': 'Hours with shortfall', 'value': results.get('hours_with_shortfall', 0)},
        {'metric': 'Number of shortfall events', 'value': results.get('n_shortfall_events', 0)},
        {'metric': 'Max event duration (h)', 'value': results.get('max_event_duration_h', 0)},
        {'metric': 'Mean event duration (h)', 'value': results.get('mean_event_duration_h', 0)},
        {'metric': 'Net CO2 emissions (Mt)', 'value': results.get('co2_emissions_Mt', float('nan'))},
        {'metric': 'Net CO2 emissions, store cross-check (Mt)', 'value': results.get('co2_emissions_storecheck_Mt', float('nan'))},
        {'metric': 'Resource adequacy electricity, Gotske basis (%)', 'value': results.get('gotske_resource_adequacy_pct', float('nan'))},
        {'metric': 'Unserved electricity (TWh)', 'value': results.get('unserved_elec_TWh', 0)},
        {'metric': 'Unserved heat (TWh)', 'value': results.get('unserved_heat_TWh', 0)},
        {'metric': 'CO2 emissions (% of 1990)', 'value': results.get('co2_emissions_pct_of_1990', float('nan'))},
        {'metric': 'Backup (emitting) generation (TWh)', 'value': results.get('backup_energy_TWh', float('nan'))},
        {'metric': 'Peak backup power (MW)', 'value': results.get('peak_backup_MW', float('nan'))},
        {'metric': 'CO2 price applied (EUR/tCO2)', 'value': results.get('co2_price_EUR_per_t', float('nan'))},
    ])
    deep.to_csv(T('deep_reliability_metrics.csv', None), index=False)
    figs['_deep'] = deep

    # summary table
    demand = results.get('total_demand_TWh', 0)
    summ = pd.DataFrame([
        {'metric': 'System cost (B EUR)', 'value': results['objective'] / 1e9},
        {'metric': 'Total served demand (TWh)', 'value': demand},
        {'metric': 'Total generation (TWh)', 'value': results.get('total_generation_TWh', 0)},
        {'metric': 'RE share (%)', 'value': results.get('re_share_pct', 0)},
        {'metric': 'Load shedding (TWh)', 'value': results.get('load_shedding_TWh', 0)},
        {'metric': 'Load shedding (% demand)',
         'value': results.get('load_shedding_TWh', 0) / demand * 100 if demand else 0},
        {'metric': 'Curtailment RE (TWh)', 'value': results.get('curtailment_TWh', 0)},
        {'metric': 'Runtime (min)', 'value': results['elapsed_seconds'] / 60}])
    summ.to_csv(T('summary.csv', None), index=False)
    figs['_summary'] = summ
    return figs


# ----------------------------------------------------------------------------
# Word helpers
# ----------------------------------------------------------------------------
def _add_table(doc, df, caption=None, max_rows=60):
    if caption:
        doc.add_paragraph().add_run(caption).bold = True
    df = df.head(max_rows)
    t = doc.add_table(rows=1, cols=len(df.columns)); t.style = 'Light Grid Accent 1'
    for j, c in enumerate(df.columns):
        t.rows[0].cells[j].text = str(c)
    for _, row in df.iterrows():
        cells = t.add_row().cells
        for j, c in enumerate(df.columns):
            v = row[c]
            cells[j].text = f"{v:,.3f}" if isinstance(v, (int, float, np.floating)) else str(v)
    doc.add_paragraph()


def _pic(doc, path, w=6.3):
    if path and os.path.exists(path):
        doc.add_picture(path, width=Inches(w))


def _pct(a, b):
    # percent change for DISPLAY. Returns NaN when the base is effectively zero
    # (a 21.8 TWh rise from 4e-7 TWh is 'from ~zero', not 5.4 billion percent)
    # or when the sign flips (a percent across a sign change is meaningless).
    if abs(a) < 1e-3:
        return float('nan')
    if (a < 0 < b) or (b < 0 < a):
        return float('nan')
    return (b - a) / a * 100


METHOD_BLURB = {
    'mixed': ("Production configuration: a per-channel assignment of the two imposition "
              "methods, each variable routed to the method that does not contaminate it. "
              "DIRECT (raw C2E future) is used for wind, solar, run-of-river and cooling, "
              "where the C2E world is kept internally coherent and its fuller stress (the "
              "capacity-factor level offset and the within-year coincidence of extremes) is "
              "what we want to see. QDM (a smoothed seasonal change factor with mean "
              "conservation; Cannon et al. 2015, Pierce et al. 2015) is used for heating "
              "demand and reservoir inflow, the two channels where DIRECT would transplant a "
              "C2E seasonal SHAPE onto Gotske infrastructure sized for a different shape, "
              "importing a dataset-definition artifact rather than a climate signal: heating, "
              "because C2E's demand profile differs from the network's own; and reservoir "
              "inflow, because C2E's river-discharge based inflow has different seasonal "
              "timing than the network's runoff based inflow, and the cyclic reservoir "
              "requires the annual volume to be conserved. Note that, because channels use "
              "different impositions, cross-variable extremes between a DIRECT channel and a "
              "QDM channel are less tightly coupled than within a single method."),
    'qdm': ("Climate-change signal isolated on the network's own weather-year chronology - "
            "PRIMARY METHOD. Wind, solar and run-of-river capacity factors are reshaped "
            "quantile by quantile by the C2E future/baseline change (Quantile Delta Mapping, "
            "Cannon et al. 2015), which is appropriate for bounded capacity factors. Heating "
            "demand and reservoir inflow, which are zero-inflated, are instead reshaped by a "
            "smoothed seasonal CHANGE FACTOR with PresRAT-style mean conservation (Pierce et "
            "al. 2015) and dry-season flooring (after Vrac et al. 2016): a per-quantile map is "
            "not meaningful where C2E and the network have different seasonal support, and for "
            "the cyclic reservoirs the annual inflow VOLUME must be conserved. Cooling is the "
            "climatological change after the embedded historical cooling is extracted. "
            "Isolates the climate-change signal: changes in variability and extremes are "
            "carried through while pipeline-level differences between C2E and the network's "
            "ERA5/atlite world cancel."),
    'direct': ("The C2E future-year world, adopted wholesale: raw substitution of the "
               "bias-corrected C2E capacity factors (SWT120_3600 turbine, the closest "
               "specific-power match to the network's assumptions), heating and hydro as "
               "energy-anchored shape transplants of the C2E future year, and the raw C2E "
               "future cooling series after extraction of the embedded historical cooling. "
               "This is Gotske's own weather-year stress philosophy extended to a future "
               "year; it additionally carries conversion-chain differences between C2E and "
               "the network's atlite world (siting rule, turbine fleet, PV model), so it "
               "stresses levels as well as the climate signal, and event coincidence "
               "(heat wave with simultaneous wind lull) is physically consistent within "
               "the C2E year."),
}


def _write_results_readme(base, method):
    txt = f"""# Results folder guide ({method}, pipeline v16)

Every table exists twice: the plain name is the CLIMATE-MODIFIED run, the
"_original" suffix is the unmodified weather-year {config.WEATHER_YEAR} reference
dispatched the same way. Compare the pair to see the climate effect.

## Folders
- 01_generation: generation mix, weekly and monthly stacks, utilisation
- 02_load_shedding: unserved energy time series, by bus, region, sector
- 03_curtailment: curtailed renewable energy by technology
- 04_regional: regional generation and balance detail
- 05_sector_coupling: cross-sector balances
- 06_storage: store levels and storage dispatch
- 07_climate_signal: the applied capacity-factor and demand changes (inputs)
- 08_reliability_deep: loss-of-load analysis, residual load duration curve
- tables: every table as csv
- networks: the dispatched networks (.nc)

## Key tables
- headline_comparison.csv: original vs modified at a glance
- adequacy_gotske.csv: electricity-only unserved energy and resource adequacy
  on the same basis as Gotske et al. (2024), heat shedding listed separately
- deep_reliability_metrics.csv: full scalar metrics including the CO2 store
  cross-check (net emissions counted two independent ways)
- cost_decomposition.csv: objective split into operating cost, net carbon
  payment, and shedding penalty. The raw objective is NOT a system cost
- seasonal_summary.csv, stress_indicators.csv: when the system is stressed
- cf_change_by_carrier.csv / cf_change_by_region.csv: the climate signal
  (these are inputs, identical for original and modified by construction)

## Reading the numbers
- Adequacy, unserved energy, shedding events are robust to the carbon price.
- Emissions, cost, and curtailment respond to the carbon price; check
  co2_price in deep_reliability_metrics.csv matches the design shadow price
  (about 468 EUR/tCO2 read from the design network file).
"""
    with open(os.path.join(base, 'README.md'), 'w') as f:
        f.write(txt)


def _headline_bullets(ro, rm):
    """Auto-generated plain-language findings for the top of the method report.
    Factual statements only; interpretation is left to the reader."""
    b = []
    try:
        adq = rm.get('gotske_resource_adequacy_pct', None)
        ue = rm.get('gotske_unserved_elec_TWh', rm.get('unserved_elec_TWh', 0.0))
        uh = rm.get('unserved_heat_TWh', 0.0)
        ev = rm.get('n_elec_events_over_24h', 0)
        if adq is not None and adq == adq:
            b.append(f"Resource adequacy (electricity, Gotske basis): {adq:.3f}%. "
                     f"Unserved electricity {ue:.3f} TWh, unserved heat {uh:.3f} TWh, "
                     f"electricity shortfall events longer than 24 h: {ev}.")
    except Exception:
        pass
    try:
        em = rm.get('co2_emissions_Mt', None)
        chk = rm.get('co2_emissions_storecheck_Mt', None)
        price = rm.get('co2_price_EUR_per_t', None)
        base = getattr(config, 'CO2_1990_BASELINE_MT', None)
        if em is not None and em == em:
            t = f"Net CO2 emissions: {em:+.1f} Mt"
            if base:
                t += f" ({em / base * 100:+.2f}% of the 1990 reference {base:.0f} Mt)"
            if price is not None:
                t += f" at a carbon price of {price:,.0f} EUR/tCO2"
            if chk is not None and chk == chk:
                t += f". Store cross-check: {chk:+.1f} Mt"
            b.append(t + ".")
    except Exception:
        pass
    try:
        do, dm = ro.get('total_demand_TWh', 0), rm.get('total_demand_TWh', 0)
        if do and dm:
            b.append(f"Served demand: {do:,.0f} TWh original to {dm:,.0f} TWh modified "
                     f"({(dm - do) / do * 100:+.1f}%).")
    except Exception:
        pass
    try:
        co, cm = ro.get('curtailment_TWh', 0), rm.get('curtailment_TWh', 0)
        b.append(f"Renewable curtailment: {co:,.0f} TWh original to {cm:,.0f} TWh modified"
                 + (f" ({(cm - co) / co * 100:+.0f}%)." if co else "."))
    except Exception:
        pass
    try:
        rso, rsm = ro.get('re_share_pct', float('nan')), rm.get('re_share_pct', float('nan'))
        if rso == rso and rsm == rsm:
            t = f"Renewable share of generation: {rso:.1f}% original, {rsm:.1f}% modified."
            bo, bm = ro.get('backup_energy_TWh', float('nan')), rm.get('backup_energy_TWh', float('nan'))
            if bo == bo and bm == bm:
                t += f" Backup (emitting) generation: {bo:,.0f} to {bm:,.0f} TWh."
            b.append(t)
    except Exception:
        pass
    return b


def build_method_report(results_orig, results_mod, method, figs_orig, figs_mod):
    base = config.run_dir(method)
    doc = Document()
    doc.add_heading(f'Climate-impact dispatch results - {method.upper()}', 0)
    if getattr(config, 'SYSTEM', 'gotske') == 'gotske':
        doc.add_paragraph(
            f"Weather year {config.WEATHER_YEAR}  |  C2E future {config.C2E_FUTURE}  |  "
            f"baseline {config.C2E_BASELINE}  |  scenario {config.SCENARIO_LABEL}").italic = True
    else:
        import systems as _sys
        doc.add_paragraph(
            f"System {config.SYSTEM} ({_sys.get(config.SYSTEM)['label']})  |  "
            f"design {config.design_id()}  |  C2E future {config.C2E_FUTURE}  |  "
            f"baseline {config.C2E_BASELINE}  |  scenario {config.SCENARIO_LABEL}").italic = True
    doc.add_paragraph(
        f"Pipeline {config.PIPELINE_VERSION}. Note: timestamps in event tables use the "
        f"network's internal snapshot calendar (its index year), which can differ from the "
        f"weather-year label; both describe the same {config.WEATHER_YEAR} meteorology."
    ).italic = True

    doc.add_heading('Headline findings', 1)
    for btxt in _headline_bullets(results_orig or {}, results_mod or {}):
        doc.add_paragraph(btxt, style='List Bullet')

    doc.add_heading('1. Method', 1)
    doc.add_paragraph(METHOD_BLURB.get(method, ''))

    doc.add_heading('1b. Channel methods and applied climate signal', 1)
    try:
        if method == getattr(config, 'RUN_PROFILE', 'mixed') and getattr(config, 'USE_CHANNEL_METHODS', False):
            _cm = dict(config.CHANNEL_METHODS)
        else:
            _cm = {k: method for k in ('supply', 'ror', 'hydro_inflow', 'heating', 'cooling', 'cop')}
        cm_df = pd.DataFrame([{'channel': k, 'method': v} for k, v in _cm.items()])
        _add_table(doc, cm_df, 'Table. Imposition method per input channel.')
    except Exception:
        pass
    _tabdir = os.path.join(base, 'tables'); _figdir = os.path.join(base, 'figures')
    os.makedirs(_figdir, exist_ok=True)
    try:
        _hyp = os.path.join(_tabdir, 'hydro_change_by_country.csv')
        if os.path.exists(_hyp):
            hydf = pd.read_csv(_hyp)
            _add_table(doc, hydf.round(3),
                       'Table. Reservoir inflow: applied annual change by country '
                       '(new/original network inflow volume; the qdm change factor '
                       'conserves the C2E annual ratio).')
            if 'applied_annual_ratio' in hydf.columns:
                try:
                    sr = hydf.set_index('country')['applied_annual_ratio'].sort_values()
                    fp = os.path.join(_figdir, 'hydro_applied_annual_ratio.png')
                    if _bar(sr, 'Reservoir inflow: applied annual future/baseline ratio',
                            'ratio (1 = unchanged)', fp):
                        _pic(doc, fp)
                except Exception as _fe:
                    doc.add_paragraph(f'(hydro ratio figure unavailable: {_fe})')
                _big = sr[(sr > 1.5) | (sr < 0.67)]
                if len(_big):
                    doc.add_paragraph(
                        'Caution: the applied annual inflow change exceeds +/-50% in: '
                        + ', '.join(f'{c} (x{v:.2f})' for c, v in _big.items())
                        + '. Each C2E year file is a SINGLE YEAR of one realization, so '
                          'these ratios mix the forced climate signal with internal '
                          'variability. Published ensemble-mean signals (e.g. Spain about '
                          '-40% in the C2E reference application, which uses CESM2 under '
                          'the harsher SSP3-7.0 over 20-year windows) are multi-decade, '
                          'multi-realization means and are NOT comparable to single-year '
                          'ratios; under the milder SSP4.5 used here the forced signal is '
                          'weaker still, so single-year ratios are even more dominated by '
                          'internal variability. Interpret single-year Southern European '
                          'hydro adequacy accordingly.')
    except Exception as e:
        doc.add_paragraph(f'(hydro signal table unavailable: {e})')
    try:
        _cvp = os.path.join(_tabdir, 'hydro_coverage.csv')
        if os.path.exists(_cvp):
            _add_table(doc, pd.read_csv(_cvp).round(2),
                       'Table. Hydro coverage. C2E covers the 10 largest reservoir '
                       'countries (80% of European inflows) and 6 run-of-river countries '
                       '(83% of generation) BY DESIGN (Wohland et al. 2025 SI Fig. S2); '
                       'uncovered units keep their design-year series.')
    except Exception:
        pass
    try:
        _htp = os.path.join(_tabdir, 'heating_change_by_country.csv')
        if os.path.exists(_htp):
            htdf = pd.read_csv(_htp)
            _add_table(doc, htdf.round(3),
                       'Table. Heating demand: qdm daily change-factor statistics and '
                       'applied annual ratio by country.')
            if 'annual_ratio' in htdf.columns:
                sr = htdf.set_index('country')['annual_ratio'].sort_values()
                fp = os.path.join(_figdir, 'heating_applied_annual_ratio.png')
                _bar(sr, 'Heating demand: applied annual future/baseline ratio',
                     'ratio (1 = unchanged)', fp)
                _pic(doc, fp)
    except Exception:
        pass
    try:
        _clp = os.path.join(_tabdir, 'cooling_change_by_country.csv')
        if os.path.exists(_clp):
            _add_table(doc, pd.read_csv(_clp).round(3),
                       'Table. Cooling demand change by country.')
    except Exception:
        pass

    doc.add_heading('2. Headline comparison (original vs climate-modified)', 1)
    if results_orig and results_mod:
        do, dm = results_orig.get('total_demand_TWh', 0), results_mod.get('total_demand_TWh', 0)
        rows = [
            ['System cost (B EUR)', results_orig['objective'] / 1e9, results_mod['objective'] / 1e9],
            ['Served demand (TWh)', do, dm],
            ['Total generation (TWh)', results_orig.get('total_generation_TWh', 0), results_mod.get('total_generation_TWh', 0)],
            ['RE share (%)', results_orig.get('re_share_pct', 0), results_mod.get('re_share_pct', 0)],
            ['Load shedding (TWh)', results_orig['load_shedding_TWh'], results_mod['load_shedding_TWh']],
            ['Curtailment RE (TWh)', results_orig['curtailment_TWh'], results_mod['curtailment_TWh']],
        ]
        df = pd.DataFrame([[m, o, mo, _pct(o, mo)] for m, o, mo in rows],
                          columns=['Metric', 'Original', 'Modified', '% change'])
        df.to_csv(os.path.join(base, 'tables', 'headline_comparison.csv'), index=False)
        _add_table(doc, df, 'Table 1. Headline metrics.')

    doc.add_heading('3. Climate signal applied (capacity factors)', 1)
    _pic(doc, figs_mod.get('cf_signal'))
    _pic(doc, figs_mod.get('cf_region') if 'cf_region' in figs_mod else None)

    doc.add_heading('4. Generation', 1)
    for k in ['genmix', 'monthly', 'weekly', 're_utilisation']:
        _pic(doc, figs_mod.get(k))

    doc.add_heading('5. Load shedding (the headline reliability metric)', 1)
    for k in ['shedding_ts', 'shedding_monthly', 'shedding_by_region', 'shedding_by_sector', 'shedding_by_bus']:
        _pic(doc, figs_mod.get(k))

    # 5b. deep reliability analysis (loss-of-load events, adequacy, emissions)
    doc.add_heading('5b. Deep reliability analysis', 1)
    if '_deep' in figs_mod:
        _add_table(doc, figs_mod['_deep'], 'Table. Deep reliability + emissions metrics (modified).')
    if '_worst_events' in figs_mod:
        _add_table(doc, figs_mod['_worst_events'], 'Table. Ten worst shortfall events (modified).')
    if '_adeq' in figs_mod:
        _add_table(doc, figs_mod['_adeq'], 'Table. Resource adequacy on the Gotske basis (modified).')
    if '_costd' in figs_mod:
        _add_table(doc, figs_mod['_costd'], 'Table. Cost decomposition of the objective (modified).')
    if '_stress' in figs_mod:
        _add_table(doc, figs_mod['_stress'], 'Table. System stress indicators (modified).')
    _pic(doc, figs_mod.get('residual_dc'))
    for k in ['lol_ts', 'lol_cum', 'lol_dc', 'shedding_event_hist', 'backup_ts']:
        _pic(doc, figs_mod.get(k))

    doc.add_heading('6. Curtailment', 1)
    _pic(doc, figs_mod.get('curtailment'))

    doc.add_heading('7. Regional detail', 1)
    for k in ['region', 'region_tech']:
        _pic(doc, figs_mod.get(k))

    doc.add_heading('8. Sector coupling and balance', 1)
    for k in ['balance']:
        _pic(doc, figs_mod.get(k))

    doc.add_heading('9. Storage', 1)
    _pic(doc, figs_mod.get('store_soc'))

    doc.add_heading('10. Data provenance and caveats', 1)
    for _t in [
        (f"Dataset and scenario. The inputs are produced with the Climate2Energy (C2E) "
         f"conversion framework (Wohland et al. 2025, arXiv:2508.09531) driven by "
         f"{config.SCENARIO_LABEL} climate simulations, per the project description and "
         f"supervisor (transient 2015-2100; the year files used here are single-year "
         f"slices: {config.C2E_BASELINE} baseline, {config.C2E_FUTURE} future). Note the "
         f"published C2E reference application uses dedicated CESM2 SSP3-7.0 runs over "
         f"1995-2015 and 2080-2100; that paper is cited for the conversion methodology, "
         f"not for these files' climate signals."),
        ("Single-year caveat. A future/baseline ratio of two single years mixes the "
         "forced climate signal with internal variability. The paper's headline changes "
         "(e.g. Southern European hydropower -10% to -40%) are 20-year, 9-realization "
         "ensemble means and are not directly comparable to single-year ratios. This is "
         "the central motivation for the multi-year sweep."),
        ("Hydro construction and why inflow uses qdm. The C2E framework converts routed "
         "river discharge to hydropower energy through a piecewise linear regression "
         "calibrated on ENTSO-E, with a spill segment above the 75th discharge "
         "percentile: run-of-river GENERATION therefore saturates in wet conditions "
         "while reservoir INFLOW (uncapped water arrival) does not, which explains calm "
         "ror ratios next to large inflow ratios in the same country. The network's own "
         "inflow is runoff-based (ERA5 via atlite) while C2E's is discharge-based, so "
         "their seasonal timing differs even for the same year; the C2E authors "
         "additionally validate their hydropower on ANNUAL totals (about 6% mean error "
         "in the reference application) while cautioning against climate-model "
         "sub-annual hydrology. The qdm inflow channel is designed for exactly this: it "
         "keeps the network's own seasonal timing (the discharge-vs-runoff definitional "
         "difference cancels in the C2E future/baseline ratio) and conserves the C2E "
         "ANNUAL change, the validated timescale, which the cyclic reservoir also "
         "physically requires. Run-of-river remains direct (bounded capacity factor, no "
         "volume constraint; saturation is part of the C2E series itself)."),
        ("Reservoir operation. Gotske's headline runs additionally constrain reservoir "
         "levels to the historical ENTSO-E minimum; that dataset is not available here, "
         "so reservoirs have full annual foresight in BOTH the original and the modified "
         "runs (documented deviation, mildly optimistic for adequacy on both sides "
         "alike)."),
        ("Comparability. The original and modified runs share the network, solver "
         "settings, carbon price and load-shedding cost; only the weather inputs differ, "
         "so all reported gaps are attributable to the imposed climate."),
    ]:
        doc.add_paragraph(_t)

    doc.add_heading('11. Summary tables', 1)
    if '_summary' in figs_mod:
        _add_table(doc, figs_mod['_summary'], 'Table 2. Modified-system summary.')

    try:
        _write_results_readme(base, method)
    except Exception:
        pass

    out = os.path.join(base, f'REPORT_{method}.docx')
    doc.save(out)
    return out


def build_comparison_report(all_results):
    cmp_dir = os.path.join(config.SESSION_DIR, f"COMPARISON_{config.design_key()}_c2e{config.C2E_FUTURE}")
    FIG = os.path.join(cmp_dir, 'figures'); TAB = os.path.join(cmp_dir, 'tables')
    os.makedirs(FIG, exist_ok=True); os.makedirs(TAB, exist_ok=True)

    doc = Document()
    doc.add_heading('Run comparison: imposing the future climate', 0)
    doc.add_paragraph(
        "The same fixed network is stress-tested under different constructions of the C2E "
        "future. MIXED is the production configuration: DIRECT (raw C2E future) for wind, "
        "solar, run-of-river and cooling, and QDM (annual-conserving change factor on the "
        "network's own chronology) for heating demand and reservoir inflow, the two channels "
        "where a raw transplant would import a dataset-definition artifact. The pure QDM and "
        "pure DIRECT runs, where present, bracket the sensitivity to that choice: pure QDM is "
        "the isolated climate signal on every channel; pure DIRECT is the C2E future wholesale, "
        "which also carries the conversion-chain level offset between C2E and the network's "
        "atlite world (visible mainly as curtailment).")

    rows = []
    for method, rr in all_results.items():
        mod, orig = rr.get('mod'), rr.get('orig')
        if not mod:
            continue
        d = mod.get('total_demand_TWh', 0)
        rows.append({'method': method, 'cost_BEUR': mod['objective'] / 1e9,
                     're_share_pct': mod.get('re_share_pct', 0),
                     'resource_adequacy_pct': mod.get('resource_adequacy_pct', float('nan')),
                     'load_shed_TWh': mod['load_shedding_TWh'],
                     'load_shed_pct': mod['load_shedding_TWh'] / d * 100 if d else 0,
                     'peak_unserved_MW': mod.get('peak_unserved_MW', float('nan')),
                     'n_shortfall_events': mod.get('n_shortfall_events', float('nan')),
                     'max_event_h': mod.get('max_event_duration_h', float('nan')),
                     'co2_emissions_Mt': mod.get('co2_emissions_Mt', float('nan')),
                     'backup_TWh': mod.get('backup_energy_TWh', float('nan')),
                     'curtailment_TWh': mod['curtailment_TWh'],
                     'cost_vs_orig_pct': _pct(orig['objective'], mod['objective']) if orig else float('nan'),
                     'unserved_elec_TWh': mod.get('gotske_unserved_elec_TWh', mod.get('unserved_elec_TWh', float('nan'))),
                     'gotske_adequacy_pct': mod.get('gotske_resource_adequacy_pct', float('nan')),
                     'demand_change_pct': _pct(orig.get('total_demand_TWh', 0), d) if orig else float('nan'),
                     'co2_price_EUR_per_t': mod.get('co2_price_EUR_per_t', float('nan')),
                     'pipeline_version': mod.get('pipeline_version', 'pre-v15'),
                     'stale_vs_current_code': mod.get('pipeline_version', None) != getattr(config, 'PIPELINE_VERSION', 'v15')})
    master = pd.DataFrame(rows)
    master.to_csv(os.path.join(TAB, 'master_metrics.csv'), index=False)
    _add_table(doc, master, 'Table 1. Modified-system metrics by method.')

    if not master.empty:
        for col, ylab, fname in [('load_shed_TWh', 'Load shedding (TWh)', 'cmp_loadshed.png'),
                                 ('curtailment_TWh', 'Curtailment RE (TWh)', 'cmp_curtailment.png'),
                                 ('cost_BEUR', 'System cost (B EUR)', 'cmp_cost.png')]:
            plt.figure(figsize=(7, 4.5))
            plt.bar(master['method'], master[col], color=[BLUE, RED, AMBER][:len(master)], edgecolor='k', lw=.3)
            plt.ylabel(ylab); plt.title(ylab + ' by method'); plt.grid(axis='y', alpha=.3)
            p = os.path.join(FIG, fname); _save(p); _pic(doc, p, 5.3)

    # per-region load shedding across methods (deep comparison table)
    reg_rows = {}
    for method, rr in all_results.items():
        mod = rr.get('mod')
        if not mod:
            continue
        for reg, v in mod.get('load_shedding_by_region_GWh', {}).items():
            reg_rows.setdefault(reg, {})[method] = v
    if reg_rows:
        reg_df = pd.DataFrame(reg_rows).T.fillna(0)
        reg_df.to_csv(os.path.join(TAB, 'load_shedding_by_region_by_method.csv'))
        doc.add_heading('2. Load shedding by region across methods (GWh)', 1)
        _add_table(doc, reg_df.reset_index().rename(columns={'index': 'region'}))

    doc.add_heading('3. Interpretation', 1)
    doc.add_paragraph(
        "MIXED is the production result: each input channel is routed to the imposition that "
        "does not contaminate it, so its heat and hydro response is the climate change on the "
        "network's own profiles while its supply side is the coherent C2E future (including "
        "the level offset, visible as curtailment). Pure QDM and pure DIRECT, where present, "
        "bracket the method sensitivity. Rows with stale_vs_current_code = 1 were produced by "
        "an OLDER pipeline version than the current one and should be re-run before being "
        "quoted next to current results; treat them as indicative only.")
    out = os.path.join(cmp_dir, f'COMPARISON_{config.design_key()}_c2e{config.C2E_FUTURE}.docx')
    doc.save(out)
    return out
